import os
import re
from docx import Document

# Function to read text files from the directory
def read_text_files(directory):
    data = {}
    for filename in os.listdir(directory):
        if filename.endswith(".txt"):
            with open(os.path.join(directory, filename), 'r') as file:
                data[filename] = file.read()
    return data

# Function to analyze MAC addresses
def analyze_mac_addresses(data):
    mac_pattern = re.compile(r"([0-9A-Fa-f]{2}[:]){5}([0-9A-Fa-f]{2})")
    valid_mac_addresses = [mac for mac in data if mac_pattern.match(mac)]
    return len(valid_mac_addresses), valid_mac_addresses

# Function to analyze passwords
def analyze_passwords(data):
    password_criteria = re.compile(r'(?=.*[A-Z])(?=.*[a-z])(?=.*\d)(?=.*[@$!%*?&])[A-Za-z\d@$!%*?&]{12,}')
    valid_passwords = [pwd for pwd in data if password_criteria.match(pwd)]
    return len(valid_passwords), valid_passwords

# Function to analyze IP addresses
def analyze_ip_addresses(data):
    ip_pattern = re.compile(r"(\d{1,3}\.){3}\d{1,3}")
    valid_ip_addresses = [ip for ip in data if ip_pattern.match(ip)]
    return len(valid_ip_addresses), valid_ip_addresses

# Function to analyze email addresses
def analyze_email_addresses(data):
    email_pattern = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")
    valid_email_addresses = [email for email in data if email_pattern.match(email)]
    return len(valid_email_addresses), valid_email_addresses

# Function to analyze bank account numbers (simple validation)
def analyze_bank_accounts(data):
    bank_pattern = re.compile(r"\d{8,12}")
    valid_bank_accounts = [account for account in data if bank_pattern.match(account)]
    return len(valid_bank_accounts), valid_bank_accounts

# Function to analyze social security numbers
def analyze_ssns(data):
    ssn_pattern = re.compile(r"\d{3}-\d{2}-\d{4}")
    valid_ssns = [ssn for ssn in data if ssn_pattern.match(ssn)]
    return len(valid_ssns), valid_ssns

# Function to analyze financial information
def analyze_financial_info(data):
    # Placeholder for complex financial info analysis
    return len(data), data

# Function to analyze health care data
def analyze_health_care_data(data):
    # Placeholder for health care data analysis
    return len(data), data

# Function to generate a Word document report
def generate_report(analysis_results):
    doc = Document()
    doc.add_heading('Data Analysis Report', 0)

    for filename, result in analysis_results.items():
        doc.add_heading(f'File: {filename}', level=1)
        doc.add_paragraph(f"Total Valid Entries: {result['count']}")
        doc.add_heading('Valid Data Samples:', level=2)
        for data in result['valid_data'][:5]:  # Show only first 5 samples
            doc.add_paragraph(data)
        doc.add_paragraph('\n')

    doc.save('Data_Analysis_Report.docx')

# Main function to handle analysis
def main():
    directory = "data_corpus"
    data_files = read_text_files(directory)

    analysis_results = {}

    for filename, content in data_files.items():
        data = content.split('\n')

        if "mac" in filename:
            count, valid_data = analyze_mac_addresses(data)
        elif "password" in filename:
            count, valid_data = analyze_passwords(data)
        elif "ip" in filename:
            count, valid_data = analyze_ip_addresses(data)
        elif "email" in filename:
            count, valid_data = analyze_email_addresses(data)
        elif "bank" in filename:
            count, valid_data = analyze_bank_accounts(data)
        elif "social" in filename:
            count, valid_data = analyze_ssns(data)
        elif "financial" in filename:
            count, valid_data = analyze_financial_info(data)
        elif "healthcare" in filename:
            count, valid_data = analyze_health_care_data(data)
        else:
            count, valid_data = 0, []

        analysis_results[filename] = {
            'count': count,
            'valid_data': valid_data
        }

    generate_report(analysis_results)
    print("Data analysis report generated: Data_Analysis_Report.docx")

if __name__ == "__main__":
    main()
